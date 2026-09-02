from __future__ import annotations

from io import BytesIO
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from reports.report_builder import AnalysisReport


def _text(value: Any) -> str:
    if value is None:
        return "Not available"
    if isinstance(value, float):
        return f"{value:g}"
    return str(value)


def _add_bullets(document: Document, values: list[str]) -> None:
    for value in values:
        document.add_paragraph(_text(value), style="List Bullet")


def _add_dataset_overview(document: Document, datasets: list[dict[str, Any]]) -> None:
    if not datasets:
        return
    document.add_heading("Dataset Overview", level=1)
    table = document.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    for cell, label in zip(table.rows[0].cells, ("Dataset", "Rows", "Columns", "Numeric columns", "Date coverage")):
        cell.text = label
    for dataset in datasets:
        row = table.add_row().cells
        temporal = "; ".join(
            f"{item.get('column')}: {item.get('min_date')} to {item.get('max_date')}"
            for item in dataset.get("datetime_coverage", [])
        )
        values = (
            dataset.get("name"), dataset.get("rows"), dataset.get("columns"),
            ", ".join(map(str, dataset.get("numeric_columns", []))), temporal,
        )
        for cell, value in zip(row, values):
            cell.text = _text(value)


def _add_findings(document: Document, findings: list[dict[str, Any]]) -> None:
    if not findings:
        return
    document.add_heading("Key Findings", level=1)
    for item in findings:
        document.add_paragraph(
            f"[{item.get('evidence_id')}] {item.get('summary')}", style="List Bullet"
        )


def _caption(visualization: dict[str, Any]) -> str:
    parts = []
    if visualization.get("description"):
        parts.append(str(visualization["description"]))
    if visualization.get("chart_type"):
        parts.append(f"Chart type: {visualization['chart_type']}.")
    if visualization.get("datasets"):
        parts.append("Dataset(s): " + ", ".join(visualization["datasets"]) + ".")
    return " ".join(parts) or "Visualization recorded during analysis."


def _add_visualizations(document: Document, visualizations: list[dict[str, Any]]) -> None:
    if not visualizations:
        return
    document.add_heading("Visualizations", level=1)
    for visualization in visualizations:
        number = visualization.get("number")
        title = visualization.get("title") or "Analysis visualization"
        document.add_heading(f"Figure {number}: {title}", level=2)
        figure = visualization.get("figure")
        embedded = False
        if figure is not None and hasattr(figure, "to_image"):
            try:
                png = figure.to_image(format="png", width=1000, height=600, scale=1)
                document.add_picture(BytesIO(png), width=Inches(6.5))
                document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                embedded = True
            except Exception:
                embedded = False
        document.add_paragraph(_caption(visualization), style="Caption")
        if not embedded:
            document.add_paragraph("The chart image could not be embedded; its recorded metadata is preserved.")


def _add_ml_results(document: Document, findings: list[dict[str, Any]]) -> None:
    ml_items = [item for item in findings if item.get("tool_name") == "train_ml_model"]
    if not ml_items:
        return
    document.add_heading("Machine Learning Results", level=1)
    for item in ml_items:
        result = item.get("result") if isinstance(item.get("result"), dict) else {}
        document.add_heading(
            f"[{item.get('evidence_id')}] {_text(result.get('task_type'))}: {_text(result.get('target_column'))}",
            level=2,
        )
        summary = document.add_table(rows=0, cols=2)
        summary.style = "Table Grid"
        split = result.get("split") if isinstance(result.get("split"), dict) else {}
        fields = (
            ("Rows used", result.get("rows_used")),
            ("Features", ", ".join(map(str, result.get("features_used", []) or []))),
            ("Evaluation strategy", split.get("strategy")),
            ("Selected model", result.get("best_model")),
            ("Selection metric", result.get("selection_metric")),
        )
        for label, value in fields:
            cells = summary.add_row().cells
            cells[0].text, cells[1].text = label, _text(value)
        models = [model for model in (result.get("models") or []) if isinstance(model, dict)][:2]
        if models:
            metrics = document.add_table(rows=1, cols=3)
            metrics.style = "Table Grid"
            metrics.rows[0].cells[0].text = "Model"
            metrics.rows[0].cells[1].text = "Role"
            metrics.rows[0].cells[2].text = "Metrics"
            for model in models:
                cells = metrics.add_row().cells
                cells[0].text = _text(model.get("name"))
                cells[1].text = "Baseline" if model.get("baseline") else "Predictive"
                values = model.get("metrics") if isinstance(model.get("metrics"), dict) else {}
                cells[2].text = "; ".join(
                    f"{name}={_text(value)}" for name, value in values.items()
                    if name not in {"confusion_matrix", "class_labels"}
                )
        associations = [item for item in (result.get("feature_associations") or []) if isinstance(item, dict)][:10]
        if associations:
            document.add_heading("Predictive Feature Associations", level=3)
            _add_bullets(document, [
                f"{entry.get('feature')}: score {_text(entry.get('association_score'))}; "
                f"relative share {_text(entry.get('relative_share_percentage'))}%"
                for entry in associations
            ])
        warnings = [str(value) for value in (result.get("warnings") or [])[:10]]
        if warnings:
            document.add_heading("ML Warnings", level=3)
            _add_bullets(document, warnings)


def render_docx(report: AnalysisReport) -> bytes:
    """Render a completed deterministic report without analysis or provider calls."""
    document = Document()
    styles = document.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)
    title = document.add_heading(report.title, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if report.objective:
        document.add_heading("Analysis Objective", level=1)
        document.add_paragraph(report.objective)
    if report.executive_summary:
        document.add_heading("Executive Summary", level=1)
        document.add_paragraph(report.executive_summary)

    _add_dataset_overview(document, report.dataset_overview)
    _add_findings(document, report.findings)
    _add_visualizations(document, report.visualizations)
    _add_ml_results(document, report.findings)

    if report.limitations:
        document.add_heading("Limitations", level=1)
        _add_bullets(document, report.limitations)
    if report.methodology:
        document.add_heading("Methodology", level=1)
        _add_bullets(document, report.methodology)
    if report.provenance:
        document.add_heading("Provenance / Evidence Appendix", level=1)
        _add_bullets(document, report.provenance)

    output = BytesIO()
    document.save(output)
    return output.getvalue()
