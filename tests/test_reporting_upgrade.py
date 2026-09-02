import json
from io import BytesIO
from types import SimpleNamespace
from unittest.mock import patch

import pandas as pd
import plotly.graph_objects as go
from docx import Document

from agent.agent import (
    Agent,
    _compact_autonomous_findings,
    _reportable_visualizations_from_findings,
)
from autonomous.results import Finding
from reports.docx_renderer import render_docx
from reports.report_builder import build_analysis_report, render_markdown


def _message(content=None, tool_calls=None):
    return SimpleNamespace(content=content, tool_calls=tool_calls)


def _tool_call(name, arguments):
    return SimpleNamespace(function=SimpleNamespace(name=name, arguments=json.dumps(arguments)))


def _datasets():
    return {
        "sales.csv": pd.DataFrame({
            "Region": ["North", "South"],
            "Sales": [10.0, 20.0],
            "Profit": [2.0, 5.0],
        }),
        "stores.csv": pd.DataFrame({"Region": ["North", "South"], "Type": ["A", "B"]}),
    }


def _artifact(title, dataset="sales.csv"):
    figure = go.Figure(go.Bar(x=["A", "B"], y=[1, 2]))
    figure.update_layout(title=title)
    return {
        "figure": figure,
        "title": title,
        "chart_type": "bar",
        "description": f"Recorded {title}.",
        "datasets": [dataset],
        "tool_name": "create_visualization",
        "finding_id": "",
    }


def test_reactive_collects_multiple_figures_in_order_and_keeps_last_top_level(sample_df):
    first = {
        "chart_type": "bar", "x_column": "Store", "y_column": "Weekly_Sales",
    }
    second = {
        "chart_type": "scatter", "x_column": "Temperature", "y_column": "Weekly_Sales",
    }
    provider = SimpleNamespace(chat=SimpleNamespace())
    provider.chat = lambda *args, **kwargs: provider.responses.pop(0)
    provider.responses = [
        _message(tool_calls=[_tool_call("create_visualization", first), _tool_call("create_visualization", second)]),
        _message("Charts created."),
        _message("Here are the charts."),
    ]
    agent = Agent()
    agent.llm = provider

    result = agent.run("Create two charts.", sample_df, autonomous=False)

    assert len(result["visualizations"]) == 2
    assert [item["sequence"] for item in result["visualizations"]] == [1, 2]
    assert result["figure"] is result["visualizations"][-1]["figure"]
    assert "figure" not in result["evidence"][0]["result"]


def test_autonomous_figure_extraction_is_ordered_bounded_and_removed_from_compaction():
    findings = []
    for index in range(7):
        artifact = _artifact(f"Chart {index + 1}")
        findings.append(Finding(
            id=f"finding_{index + 1}", step_id=f"step_{index + 1}",
            tool_name="create_visualization", datasets=["sales.csv"], result=artifact,
        ))

    visualizations = _reportable_visualizations_from_findings(findings)
    compact = _compact_autonomous_findings(findings)

    assert len(visualizations) == 5
    assert [item["title"] for item in visualizations] == [f"Chart {i}" for i in range(1, 6)]
    assert all("figure" not in item["result"] for item in compact)


def test_report_supports_no_figure_legacy_figure_and_multiple_figure_markdown():
    base = {"answer": "Complete.", "evidence": [], "trace": [], "figure": None}
    assert build_analysis_report("Analyze", base, _datasets()).visualizations == []

    legacy = dict(base, figure=_artifact("Legacy")["figure"])
    assert len(build_analysis_report("Analyze", legacy, _datasets()).visualizations) == 1

    result = dict(base, visualizations=[_artifact("First"), _artifact("Second", "stores.csv")])
    markdown = render_markdown(build_analysis_report("Analyze", result, _datasets()))
    assert markdown.index("Figure 1: First") < markdown.index("Figure 2: Second")
    assert "Dataset(s): stores.csv." in markdown
    assert "data:image" not in markdown


def test_docx_contains_expected_sections_dataset_table_and_embedded_figures():
    evidence = [{"tool_name": "statistics", "result": {"column": "Sales", "mean": 15.0}}]
    result = {
        "answer": "Average sales are 15.", "trace": [], "evidence": evidence,
        "figure": None, "visualizations": [_artifact("Sales chart"), _artifact("Profit chart")],
    }
    report = build_analysis_report("Summarize performance", result, _datasets())

    with patch("agent.agent.Agent.run", side_effect=AssertionError("must not run analysis")):
        payload = render_docx(report)
    document = Document(BytesIO(payload))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)

    for heading in (
        "AI Data Analysis Report", "Analysis Objective", "Executive Summary",
        "Dataset Overview", "Key Findings", "Visualizations", "Methodology",
        "Provenance / Evidence Appendix",
    ):
        assert heading in text
    assert len(document.tables) >= 1
    assert len(document.inline_shapes) == 2
    assert "Figure 1: Sales chart" in text and "Figure 2: Profit chart" in text


def test_docx_figure_failure_is_isolated_and_keeps_caption():
    class BrokenFigure:
        def to_image(self, **kwargs):
            raise RuntimeError("renderer unavailable")

    result = {
        "answer": "Complete.", "trace": [], "evidence": [], "figure": None,
        "visualizations": [{
            "figure": BrokenFigure(), "title": "Unavailable chart", "chart_type": "line",
            "description": "Monthly trend.", "datasets": ["sales.csv"],
        }, _artifact("Working chart")],
    }
    document = Document(BytesIO(render_docx(build_analysis_report("Analyze", result, _datasets()))))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)

    assert "The chart image could not be embedded" in text
    assert "Monthly trend." in text
    assert len(document.inline_shapes) == 1


def test_docx_renders_recorded_ml_evidence_and_multi_dataset_provenance():
    ml_result = {
        "task_type": "regression", "target_column": "Sales", "rows_used": 100,
        "features_used": ["Profit"], "best_model": "linear_regression", "selection_metric": "rmse",
        "split": {"strategy": "random", "train_rows": 80, "test_rows": 20},
        "models": [
            {"name": "mean_baseline", "baseline": True, "metrics": {"rmse": 4.0}},
            {"name": "linear_regression", "baseline": False, "metrics": {"rmse": 2.0, "r2": 0.7}},
        ],
        "feature_associations": [{"feature": "Profit", "association_score": 0.8, "relative_share_percentage": 100.0}],
        "warnings": ["Held-out estimate only."],
    }
    finding = Finding(
        id="finding_1", step_id="ml", tool_name="train_ml_model",
        datasets=["sales.csv", "stores.csv"], result=ml_result,
    )
    report = build_analysis_report(
        "Predict sales", {"answer": "Model evaluated.", "findings": [finding], "trace": [], "figure": None},
        _datasets(),
    )
    document = Document(BytesIO(render_docx(report)))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)

    assert "Machine Learning Results" in text
    assert "linear_regression" in text
    assert "Profit" in text
    assert any("sales.csv" in line and "stores.csv" in line for line in report.provenance)
